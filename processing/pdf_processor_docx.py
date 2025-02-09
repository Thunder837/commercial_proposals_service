import os
import re
import unidecode
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches
from docx2pdf import convert as convert_docx_to_pdf

def convert_pdf_to_docx(input_pdf, temp_docx):
    """
    Конвертирует PDF в DOCX с помощью библиотеки pdf2docx.
    """
    cv = Converter(input_pdf)
    cv.convert(temp_docx, start=0, end=None)
    cv.close()
    print("PDF успешно конвертирован в DOCX:", temp_docx)

def process_docx(temp_docx, processed_docx, logo_path):
    """
    Обрабатывает DOCX:
      - Удаляет контакты подрядчика (по шаблонам для "Телефон:" и "Email:").
      - Находит цену в формате "Цена: число" и заменяет число на значение с наценкой +30%.
      - Добавляет логотип компании в начало документа.
      - Добавляет наши контактные данные в конец документа.
    """
    doc = Document(temp_docx)
    
    # Примерные регулярные выражения для удаления контактов и замены цены
    phone_pattern = r'(Телефон:\s*\+?\d[\d\s()-]{7,}\d)'
    email_pattern = r'(Email:\s*\S+@\S+)'
    price_pattern = r'(Цена:\s*)(\d+(?:[.,]\d+)?)'
    
    # Выводим для отладки содержимое документа до обработки
    print("Содержимое документа до обработки:")
    for i, para in enumerate(doc.paragraphs):
        print(f"Параграф {i}: {para.text}")
    
    for para in doc.paragraphs:
        text = para.text
        text = re.sub(phone_pattern, '', text)
        text = re.sub(email_pattern, '', text)
        
        # Функция для замены цены с наценкой +30%
        def replace_price(match):
            prefix = match.group(1)
            price_str = match.group(2).replace(',', '.')
            try:
                price = float(price_str)
                new_price = price * 1.30
                return prefix + f"{new_price:.2f}"
            except Exception:
                return match.group(0)
        
        text = re.sub(price_pattern, replace_price, text)
        para.clear()
        para.add_run(text)
    
    # Добавляем логотип компании в начало документа
    try:
        first_paragraph = doc.paragraphs[0]
        run = first_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(2))
        print("Логотип успешно добавлен.")
    except Exception as e:
        print("Ошибка при добавлении логотипа:", e)
    
    # Добавляем наши контактные данные в конец документа
    doc.add_paragraph("\nКонтакты: info@gorbachevmedia.com, +79771463998")
    
    doc.save(processed_docx)
    print("Обработка DOCX завершена. Сохранён файл:", processed_docx)

def convert_docx_to_pdf_file(processed_docx, output_pdf):
    """
    Конвертирует обработанный DOCX в PDF с помощью docx2pdf.
    """
    convert_docx_to_pdf(processed_docx)
    temp_pdf = os.path.splitext(processed_docx)[0] + ".pdf"
    if os.path.exists(temp_pdf):
        os.replace(temp_pdf, output_pdf)
        print("Конвертация DOCX в PDF завершена. Итоговый файл:", output_pdf)
    else:
        raise Exception("Конвертация в PDF не удалась.")

def process_pdf_via_docx(input_pdf, output_pdf, logo_path):
    """
    Полная обработка:
      1. Конвертирует исходный PDF в DOCX.
      2. Обрабатывает DOCX (удаляет контакты подрядчика, заменяет цену, добавляет логотип и новые контакты).
      3. Конвертирует обработанный DOCX обратно в PDF.
    Итоговый файл сохраняется в формате PDF.
    """
    temp_docx = "temp_processed.docx"
    processed_docx = "processed.docx"
    
    convert_pdf_to_docx(input_pdf, temp_docx)
    process_docx(temp_docx, processed_docx, logo_path)
    convert_docx_to_pdf_file(processed_docx, output_pdf)
    
    # Удаляем временные файлы
    for file in [temp_docx, processed_docx]:
        if os.path.exists(file):
            os.remove(file)
            print("Удалён временный файл:", file)
